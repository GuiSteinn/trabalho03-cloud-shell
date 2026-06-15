import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).with_name("Relatorio_Final_Trabalho_03.docx")
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
BORDER = "DADCE0"


def set_run_font(run, size=11, bold=False, color=BLACK):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), BORDER)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    return paragraph


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(11)
normal.font.color.rgb = BLACK
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

heading_tokens = {
    "Heading 1": (20, 20, 6, BLACK),
    "Heading 2": (16, 18, 6, BLACK),
    "Heading 3": (14, 16, 4, RGBColor(67, 67, 67)),
}
for style_name, (size, before, after, color) in heading_tokens.items():
    style = doc.styles[style_name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = False
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_style in ("List Bullet", "List Number"):
    style = doc.styles[list_style]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

# Titulo simples para importacao no Google Docs; nao usa o estilo Word Title.
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(3)
run = title.add_run("Relatório Final - Trabalho 03")
set_run_font(run, size=26, bold=False)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(18)
run = subtitle.add_run("Linux, Shell Script e Automação Operacional aplicada à Cloud")
set_run_font(run, size=13, color=MUTED)

metadata = [
    ("Aluno", "Guilherme Stein Zunino Sgrott"),
    ("Instituição", "Unidavi"),
    ("Tema", "Sistema de Reservas para um Hotel"),
    ("Data", "15 de junho de 2026"),
]
for label, value in metadata:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    first = paragraph.add_run(f"{label}: ")
    set_run_font(first, bold=True)
    second = paragraph.add_run(value)
    set_run_font(second)

add_heading(doc, "1. Introdução", 1)
add_body(doc, "Este trabalho apresenta um ambiente Linux containerizado para apoiar a operação de um Sistema de Reservas para um Hotel. A proposta evolui o cenário dos Trabalhos 01 e 02 para uma etapa prática de DevOps, automatizando tarefas recorrentes de administração, publicação, segurança, backup e monitoramento.")
add_body(doc, "O ambiente foi organizado para permitir que o professor construa a imagem, inicie o container, execute os scripts e acesse o portal seguindo apenas as instruções do README.md.")

add_heading(doc, "2. Objetivos", 1)
add_bullet(doc, "Executar um servidor Ubuntu isolado com Docker e Docker Compose.")
add_bullet(doc, "Publicar no Apache um site estático relacionado às reservas do hotel.")
add_bullet(doc, "Automatizar atualização, estrutura, deploy, backup e relatório com Bash.")
add_bullet(doc, "Monitorar CPU, memória, disco e disponibilidade do Apache.")
add_bullet(doc, "Aplicar usuários, grupos e permissões seguindo o menor privilégio.")

add_heading(doc, "3. Arquitetura do Ambiente", 1)
add_body(doc, "A imagem utiliza Ubuntu 24.04 e instala Apache, curl, procps, tar e bc. O container recebe o nome trabalho03-linux e publica a porta 80 na porta 8080 do computador. O entrypoint cria a estrutura temática, realiza o deploy e mantém o Apache em primeiro plano.")
add_body(doc, "Fluxo técnico: Navegador → localhost:8080 → container Ubuntu → Apache → /var/www/html.")
add_body(doc, "Persistência: o volume trabalho03_hotel_dados armazena dados operacionais. As pastas backups/ e logs/ são mapeadas para o computador, permitindo consulta mesmo após a remoção do container.")

add_heading(doc, "4. Estrutura Temática", 1)
add_body(doc, "O script de estrutura cria diretórios específicos para a operação hoteleira:")
for path in ("/app/hotel/reservas", "/app/hotel/hospedes", "/app/hotel/quartos", "/app/hotel/dados", "/app/hotel/logs", "/app/hotel/publicacao", "/app/hotel/backups"):
    add_bullet(doc, path)
add_body(doc, "Também são criados arquivos iniciais para reservas pendentes, hóspedes ativos, status de quartos e configuração do ambiente.")

add_heading(doc, "5. Scripts Desenvolvidos", 1)
rows = [
    ("01_update.sh", "Atualiza índices e pacotes do Ubuntu em modo não interativo."),
    ("02_apache.sh", "Instala, inicia, valida e exibe a versão do Apache."),
    ("03_estrutura.sh", "Cria a estrutura temática e remove somente temporários autorizados."),
    ("04_backup.sh", "Gera backup .tar.gz do ambiente e do site com data e hora."),
    ("05_deploy.sh", "Limpa o DocumentRoot, publica source/ e valida index.html."),
    ("06_processos.sh", "Lista, busca e encerra processos por PID com bloqueios de segurança."),
    ("07_monitoramento.sh", "Coleta CPU, RAM, disco e status do Apache com alertas."),
    ("08_usuarios_permissoes.sh", "Cria hotel_ops e reservas_user e aplica chown/chmod."),
    ("09_relatorio.sh", "Consolida informações em logs/relatorio_execucao.txt."),
    ("menu.sh", "Integra as rotinas em um menu interativo."),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
set_table_geometry(table, [2520, 6840])
set_table_borders(table)
headers = ("Script", "Finalidade")
for index, text in enumerate(headers):
    cell = table.rows[0].cells[index]
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    set_run_font(run, bold=True)
mark_header_row(table.rows[0])
for script, description in rows:
    cells = table.add_row().cells
    for index, text in enumerate((script, description)):
        cells[index].text = ""
        run = cells[index].paragraphs[0].add_run(text)
        set_run_font(run)
set_table_geometry(table, [2520, 6840])

add_heading(doc, "6. Segurança e Permissões", 1)
add_body(doc, "As rotinas que alteram pacotes, usuários ou proprietários validam a execução como root. O usuário reservas_user é uma conta de sistema sem shell de login, e o grupo hotel_ops concentra o acesso operacional.")
add_body(doc, "Os diretórios usam os modos 2750 e 2770. O bit setgid mantém a herança do grupo hotel_ops, e o projeto não utiliza chmod 777. A limpeza do deploy é permitida somente em /var/www/html, enquanto o gerenciamento de processos impede chamadas sem PID e bloqueia PIDs críticos.")

add_heading(doc, "7. Deploy, Backup e Monitoramento", 1)
add_heading(doc, "7.1 Deploy", 2)
add_body(doc, "O site estático foi adaptado do contexto funcional do Trabalho 02. O script copia index.html, sobre.html e assets/ para o Apache, configura proprietário www-data e valida a existência do arquivo principal.")
add_heading(doc, "7.2 Backup", 2)
add_body(doc, "O backup utiliza tar com gzip e recebe o padrão backup_hotel_reservas_AAAA-MM-DD_HH-MM-SS.tar.gz. O script verifica se o arquivo foi criado e possui conteúdo antes de registrar sucesso.")
add_heading(doc, "7.3 Monitoramento", 2)
add_body(doc, "A CPU é calculada a partir de /proc/stat; memória e disco são obtidos com free e df. O limite padrão é 80%, podendo ser modificado por variáveis de ambiente para demonstrar alertas. O status do Apache é consultado por pgrep.")

add_heading(doc, "8. Validação Técnica", 1)
add_body(doc, "Foram validados em container temporário: sintaxe Bash dos scripts, criação da estrutura, versão e processo do Apache, deploy, criação do backup, listagem e busca de processos, monitoramento, usuários, permissões, relatório e resposta HTTP contendo o portal Hotel Aurora.")
add_body(doc, "A imagem Docker foi construída com sucesso. Na máquina de desenvolvimento, a porta 8080 já estava ocupada por outro projeto; por isso o teste final do Trabalho 03 foi iniciado na porta alternativa 8081 por meio da variável APACHE_PORT, mantendo 8080 como padrão da entrega.")
add_body(doc, "A execução do script de atualização deve ser registrada novamente na máquina do aluno para produzir a evidência específica solicitada, pois a comunicação com o Docker Desktop ficou indisponível durante essa etapa de teste.")

add_heading(doc, "9. Evidências", 1)
add_body(doc, "A pasta evidencias/ contém um roteiro com os comandos e nomes sugeridos para as capturas. Antes da entrega, devem ser incluídas imagens que comprovem:")
for item in (
    "container e volume Docker em execução;",
    "permissão de execução dos scripts;",
    "atualização do Ubuntu e validação do Apache;",
    "estrutura temática, backup e deploy;",
    "site acessível no navegador;",
    "monitoramento, usuários, permissões e relatório;",
    "imagem publicada no DockerHub.",
):
    add_bullet(doc, item)

add_heading(doc, "10. Dificuldades Encontradas", 1)
add_body(doc, "As principais dificuldades técnicas foram adaptar o gerenciamento do Apache a um container sem systemd, proteger operações de remoção e encerramento de processos, manter persistência entre recriações do container e conciliar permissões restritas com a operação do servidor web.")
add_body(doc, "Também foi necessário tratar conflito de porta no computador de teste. A solução foi parametrizar a porta externa com APACHE_PORT, preservando o valor 8080 exigido no enunciado.")

add_heading(doc, "11. Uso de Inteligência Artificial", 1)
add_body(doc, "Foi utilizada a ferramenta ChatGPT/Codex como apoio na estruturação inicial dos scripts, revisão dos comandos, organização do README, criação do site estático e identificação de validações de segurança. A ferramenta também auxiliou na execução dos testes e na documentação dos resultados.")
add_body(doc, "O conteúdo deve ser revisado pelo aluno antes da entrega. O aprendizado envolveu funções e parâmetros em Bash, códigos de saída, Docker Compose, Apache em containers, tar, ps, pgrep, kill, /proc, free, df, chown, chmod, grupos e usuários de sistema.")

add_heading(doc, "12. Links da Entrega", 1)
add_body(doc, "GitHub: adicionar o link do repositório após a publicação.")
add_body(doc, "DockerHub: https://hub.docker.com/r/guilhermesteinn/trabalho03-hotel-shell")

add_heading(doc, "13. Conclusão", 1)
add_body(doc, "O projeto demonstra como rotinas repetitivas de administração Linux podem ser padronizadas com Shell Script e Docker. A solução integra infraestrutura, publicação, segurança, monitoramento e documentação em um ambiente reproduzível e diretamente relacionado ao Sistema de Reservas para um Hotel.")

doc.core_properties.title = "Relatório Final - Trabalho 03"
doc.core_properties.subject = "Linux, Shell Script e Automação Operacional aplicada à Cloud"
doc.core_properties.author = "Guilherme Stein Zunino Sgrott"
doc.core_properties.keywords = "Linux, Shell Script, Docker, Apache, Cloud, Hotel"
doc.save(OUTPUT)
print(OUTPUT)
